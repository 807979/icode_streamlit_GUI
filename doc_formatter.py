from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_ALIGN_VERTICAL
import re

def add_page_number(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def add_table_of_contents(doc):
    """Add a Table of Contents field to the document."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    # Add TOC field
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.text = "TOC \\o \"1-4\" \\h \\z \\u"  # TOC field code

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def create_formatted_doc(input_docx: str, filename="COBOL_Documentation9.docx", logo_path="logo.png"):
   
    input_doc = Document(input_docx)

    doc = Document()
    # === Cover Page ===
    cover_page = doc.add_paragraph()
    cover_run = cover_page.add_run("COBOL Program Reference Guide")
    cover_run.bold = True
    cover_run.font.size = Pt(24) 
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Created via LLM API").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # === Table of Contents Placeholder ===
    doc.add_paragraph("Table of Contents", style='Heading 1')
    #doc.add_paragraph("<<TOC will be updated in Word manually>>")
    add_table_of_contents(doc)
    doc.add_page_break()

    # === Header with Logo and Styled Text ===
    section = doc.sections[0]
    header = section.header
    section.top_margin = Inches(2.0)  # Adjust the value as needed

    table = header.add_table(rows=1, cols=2,width=Inches(6))
    table.allow_autofit = True
    table.autofit = True

    cell_text = table.cell(0, 0)
    para = cell_text.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT


  # Add styled text 
    run_text = para.add_run("  COBOL Program Reference Guide")
    run_text.font.name = 'Arial'
    run_text.font.size = Pt(8)
    run_text.font.color.rgb = RGBColor(255, 255, 255)
    
    # Try to add logo
    try:
        
        cell_logo = table.cell(0, 1)
        para1=cell_logo.paragraphs[0]
        para1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = para1.add_run()
        run.add_picture(logo_path, width=Inches(1.0))
      
    except Exception as e:
        #print(f"Error adding logo: {e}")
        
        para1.add_run("[Logo Missing]")

  
    
        
    # Set header background color to blue
    for cell in table.row_cells(0):
        cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell=table.cell(0, 0)  
    for cell in table.row_cells(0):
        tc_pr = cell._tc.get_or_add_tcPr()
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:val'), 'clear')
        shading_elm.set(qn('w:color'), 'auto')
        shading_elm.set(qn('w:fill'), '007ACC')  # Blue
        tc_pr.append(shading_elm)
   
  
    # === Footer with Page Number ===
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_paragraph.text = "Page "
    add_page_number(footer_paragraph)

    # === Append Original Content ===
    for element in input_doc.element.body:
        doc.element.body.append(element)

    doc.save(filename)
    print(f" Document saved as: {filename}")



#create_formatted_doc("finala.docx", logo_path="logo.png", filename="styled_output.docx")
